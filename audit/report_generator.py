import os
import datetime
import json
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from .utils_text import parse_json_insight_to_table
from .utils_analysis import wasted_spend_analyzer
from .utils_text import clean, safe_parse_gemini_json


def add_industry_benchmark_overlay(df, benchmarks):
    for metric in ["CTR", "Avg CPC", "CPA ($)", "CVR"]:
        if metric in df.columns and metric in benchmarks:
            df[f"{metric} Δ"] = df[metric] - benchmarks[metric]
            df[f"{metric} Flag"] = df.apply(
                lambda x: 'Above' if x[f"{metric} Δ"] > 0.2 * benchmarks[metric]
                else 'Below' if x[f"{metric} Δ"] < -0.2 * benchmarks[metric]
                else '',
                axis=1
            )
    return df


def generate_report(df_30, kw_df, hour_pivot, hour_raw_df,
                    insight_30, insight_kw, insight_hour,
                    geo_df, insight_geo,
                    wasted_flags, wasted_insight,
                    lp_audit_rows,
                    risk_opp_insights, lp_flags=None, competitor_insights=None):

    doc = Document()
    doc.add_heading("Google Ads Audit Report", 0)

    def autofit_table(table):
        table.allow_autofit = True
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    def add_table(title, df, columns):
        if df is None or df.empty:
            return
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        autofit_table(table)
        hdr = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr[i].text = col
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(columns):
                val = row.get(col, "")
                if col in ["Avg CPC", "CPA ($)", "Cost ($)"]:
                    try:
                        val = f"${float(val):.2f}" if val else "$0.00"
                    except Exception:
                        val = str(val)
                elif col == "CTR":
                    try:
                        val = f"{float(val) * 100:.2f}%" if val else "0.00%"
                    except Exception:
                        pass
                cells[i].text = str(val)

    def add_json_insight_section(title, json_text):
        doc.add_heading(title, level=1)
        df = parse_json_insight_to_table(json_text)
        if df is not None and not df.empty:
            add_table(title, df, df.columns.tolist())
        else:
            doc.add_paragraph("⚠️ Unable to parse structured insights — showing raw output below.")
            if isinstance(json_text, str):
                doc.add_paragraph(json_text.strip())
            else:
                doc.add_paragraph(str(json_text))

    def add_hourly_pivot(pivot):
        doc.add_heading("Hourly Performance Pivot", level=1)
        try:
            day_order = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
            pivot = pivot.reindex(day_order)
        except Exception:
            pass
        try:
            for metric in pivot.columns.levels[0]:
                sub_df = pivot[metric].replace("", 0)
                sub_df = sub_df.loc[(sub_df != 0).any(axis=1), (sub_df != 0).any(axis=0)]
                if sub_df.empty:
                    continue
                doc.add_paragraph(f"{metric}")
                table = doc.add_table(rows=sub_df.shape[0] + 1, cols=sub_df.shape[1] + 1)
                table.style = 'Table Grid'
                autofit_table(table)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Day/Hour"
                for j, col in enumerate(sub_df.columns):
                    hdr_cells[j + 1].text = str(col)
                for i, idx in enumerate(sub_df.index):
                    row_cells = table.rows[i + 1].cells
                    row_cells[0].text = str(idx)
                    for j, col in enumerate(sub_df.columns):
                        val = sub_df.loc[idx, col]
                        row_cells[j + 1].text = f"{val:.2f}" if isinstance(val, (int, float)) and val != 0 else ""
        except Exception:
            pass

    def add_heatmaps():
        for metric in ["Clicks", "Conversions", "CVR"]:
            img_path = f"report_images/{metric}_heatmap.png"
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6))
                except Exception:
                    pass

    # --- NEW: Risks and Opportunities split ---
    def add_risks_opportunities():
        try:
            data = risk_opp_insights
            if isinstance(data, str):
                data = json.loads(data)
            if not isinstance(data, dict):
                raise ValueError("Invalid risk_opp_insights format")

            risks = pd.DataFrame(data.get("Risks", []))
            opps = pd.DataFrame(data.get("Opportunities", []))

            # Risks section
            doc.add_heading("⚠️ Risks", level=1)
            if not risks.empty:
                add_table("Risks", risks, risks.columns.tolist())
            else:
                doc.add_paragraph("No Risk insights generated.")

            # Opportunities section
            doc.add_heading("✅ Opportunities", level=1)
            if not opps.empty:
                add_table("Opportunities", opps, opps.columns.tolist())
            else:
                doc.add_paragraph("No Opportunity insights generated.")

        except Exception as e:
            doc.add_heading("⚠️ Risks", level=1)
            doc.add_paragraph(f"Failed to parse Risks: {e}")
            doc.add_paragraph(str(risk_opp_insights))

            doc.add_heading("✅ Opportunities", level=1)
            doc.add_paragraph("Parsing failed.")
    # --- END NEW ---

    # Benchmarks
    benchmarks = {"CTR": 0.03, "CVR": 0.05, "Avg CPC": 2.0, "CPA ($)": 20.0}
    df_30_bench = add_industry_benchmark_overlay(df_30.copy() if df_30 is not None else pd.DataFrame(), benchmarks)

    # --- Campaign Section ---
    add_table("Campaign Performance (w/ Benchmark Overlay)", df_30_bench, [
        "Campaign Name", "CTR", "Cost ($)", "Clicks", "Impressions",
        "CTR Flag", "Avg CPC", "Avg CPC Flag", "CPA ($)", "CPA ($) Flag"
    ])
    add_json_insight_section("Campaign Insights", insight_30)

    # --- Keyword Section ---
    add_table("Keyword Performance", kw_df, [
        "Ad Group", "Keyword", "Match Type", "Quality Score", "Impressions",
        "Clicks", "CTR", "Avg CPC", "CPA ($)"
    ])
    add_json_insight_section("Keyword Insights", insight_kw)

    # --- Wasted Spend ---
    if wasted_flags:
        if wasted_insight and wasted_insight.strip():
            add_json_insight_section("Wasted Spend Insights", wasted_insight)
        else:
            doc.add_heading("Wasted Spend Insights", level=1)
            doc.add_paragraph("No insights generated by Gemini.")

    # --- Landing Page Audit Section ---
    if lp_audit_rows:
        doc.add_heading("Landing Page Audit Insights", level=1)
        for raw_json in lp_audit_rows:
            try:
                data = json.loads(raw_json)
                if isinstance(data, dict):
                    data = [data]
                df = pd.DataFrame(data)
                if df.empty:
                    raise ValueError("Empty LP audit data")
                cols = ["URL"] + [c for c in df.columns if c != "URL"]
                url = df["URL"].iloc[0]
                doc.add_heading(f"Landing Page: {url}", level=2)
                add_table("Landing Page Insights", df, cols)
            except Exception:
                doc.add_paragraph("⚠️ Failed to parse LP audit JSON — showing raw output.")
                doc.add_paragraph(raw_json.strip() if raw_json else "")

    # --- Geo Section ---
    add_table("Geographical Performance", geo_df, [
        "City", "Region", "Country", "Type", "Impressions", "Clicks",
        "Conversions", "Cost ($)", "CVR", "CPA ($)"
    ])
    add_json_insight_section("Geographical Insights", insight_geo)

    # --- Hourly Section ---
    add_hourly_pivot(hour_pivot)
    add_json_insight_section("Hourly Patterns Insights", insight_hour)
    add_heatmaps()

    # --- Competitor Intelligence ---
    if competitor_insights is not None:
        try:
            if isinstance(competitor_insights, str):
                parsed = safe_parse_gemini_json(competitor_insights)
                competitor_insights = pd.DataFrame(parsed)
            elif isinstance(competitor_insights, list):
                competitor_insights = pd.DataFrame(competitor_insights)
        except Exception as e:
            print(f"⚠️ Competitor insights parse error: {e}")
            competitor_insights = pd.DataFrame()
        if isinstance(competitor_insights, pd.DataFrame) and not competitor_insights.empty:
            doc.add_heading("Competitor Intelligence", level=1)
            add_table(
                "Top Competitor Insights",
                competitor_insights,
                ["Competitor", "Strengths", "Recommendations"]
            )
        else:
            doc.add_heading("Competitor Intelligence", level=1)
            doc.add_paragraph("No competitor insights found or could be parsed.")

    # --- Risks & Opportunities (separate tabs) ---
    add_risks_opportunities()

    filename = f"google_ads_audit_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    print(f"✅ Report saved as {filename}")
    return os.path.abspath(filename)
