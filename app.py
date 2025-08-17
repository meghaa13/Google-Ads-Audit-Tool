import os
import re
import json
import yaml
import ast
from copy import deepcopy
from flask import Flask, render_template, request, redirect, url_for, send_from_directory, session
from flask_session import Session
from google_auth_oauthlib.flow import Flow
from google.ads.googleads.client import GoogleAdsClient
from google.oauth2 import id_token
from google.auth.transport import requests
from dotenv import load_dotenv

# your report logic
from audit.main_runner import generate_google_ads_report

# ====== Basic app config ====================================================
load_dotenv()
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'  # dev only (use HTTPS in prod)

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = "generated_reports"
app.config["SESSION_TYPE"] = "filesystem"
app.secret_key = os.environ.get("FLASK_SECRET", "mushy_baby_love")
app.config["SESSION_FILE_DIR"] = os.path.abspath("flask_session")
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_USE_SIGNER"] = True
Session(app)

os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs("user_tokens", exist_ok=True)

# ----- Base google-ads YAML -------------------------------------------------
BASE_YAML_PATH = "base_google-ads.yaml"

def _load_base_config_from_file(path: str) -> dict:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    except Exception:
        return {}

if os.path.isfile(BASE_YAML_PATH):
    base_config = _load_base_config_from_file(BASE_YAML_PATH)
else:
    base_config = {
        "developer_token": os.getenv("GOOGLE_ADS_DEVELOPER_TOKEN"),
        "client_id": os.getenv("GOOGLE_ADS_CLIENT_ID"),
        "client_secret": os.getenv("GOOGLE_ADS_CLIENT_SECRET"),
        "refresh_token": os.getenv("GOOGLE_ADS_REFRESH_TOKEN") or None,
        "login_customer_id": os.getenv("GOOGLE_ADS_LOGIN_CUSTOMER_ID") or None,
    }

_required = ["developer_token", "client_id", "client_secret"]
_missing = [k for k in _required if not base_config.get(k)]
if _missing:
    raise RuntimeError(
        "Missing Google Ads base configuration. Provide 'base_google-ads.yaml' or set env vars: "
        "GOOGLE_ADS_DEVELOPER_TOKEN, GOOGLE_ADS_CLIENT_ID, GOOGLE_ADS_CLIENT_SECRET. "
        f"Missing: {', '.join(_missing)}"
    )

# ----- Fixed scopes ---------------------------------------------------------
SCOPES = [
    "https://www.googleapis.com/auth/adwords",
    "openid",
    "https://www.googleapis.com/auth/userinfo.email"
]
REDIRECT_URI = os.environ.get("REDIRECT_URI", "http://localhost:5000/callback")

# ====== OAuth Flow helper ===================================================
def get_flow(state=None):
    """
    Build a Flow either from client-secrets-web.json (if present)
    or from environment variables (Render deployment).
    """
    if os.path.isfile("client-secrets-web.json"):
        # Local dev mode: load from JSON file
        return Flow.from_client_secrets_file(
            "client-secrets-web.json",
            scopes=SCOPES,
            redirect_uri=REDIRECT_URI,
            state=state
        )

    # Production mode: build config from environment variables
    client_id = os.getenv("GOOGLE_ADS_CLIENT_ID")
    client_secret = os.getenv("GOOGLE_ADS_CLIENT_SECRET")
    #js_origin = os.getenv("JS_ORIGIN", "http://localhost:5000")

    if not client_id or not client_secret:
        raise RuntimeError("Missing OAuth credentials: set GOOGLE_ADS_CLIENT_ID and GOOGLE_ADS_CLIENT_SECRET")

    client_config = {
        "web": {
            "client_id": client_id,
            "client_secret": client_secret,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [REDIRECT_URI],
            #"javascript_origins": [js_origin],
        }
    }
    return Flow.from_client_config(
        client_config,
        scopes=SCOPES,
        redirect_uri=REDIRECT_URI,
        state=state
    )

# ====== Helpers =============================================================
def user_yaml_path_for_email(email: str) -> str:
    return os.path.join("user_tokens", f"{email.lower()}.yaml")

def load_persisted_users():
    users = {}
    for fname in os.listdir("user_tokens"):
        if not (fname.endswith(".yaml") or fname.endswith(".yml")):
            continue
        path = os.path.join("user_tokens", fname)
        key = os.path.splitext(fname)[0]
        users[key] = path
    return users

def read_yaml(path):
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f) or {}

def write_yaml(path, data):
    with open(path, "w", encoding="utf-8") as f:
        yaml.safe_dump(data, f, sort_keys=False)
    try:
        os.chmod(path, 0o600)
    except Exception:
        pass

def normalize_customer_id(cid: str) -> str:
    return re.sub(r"[^0-9]", "", cid or "")

def load_client_with_optional_login(auth_file: str, login_customer_id: str | None):
    cfg = read_yaml(auth_file)
    cfg = cfg or {}
    cfg.setdefault("developer_token", base_config.get("developer_token"))
    cfg.setdefault("client_id", base_config.get("client_id"))
    cfg.setdefault("client_secret", base_config.get("client_secret"))
    cfg.setdefault("refresh_token", base_config.get("refresh_token"))
    if login_customer_id:
        cfg["login_customer_id"] = normalize_customer_id(login_customer_id)
    return GoogleAdsClient.load_from_dict(cfg)

PERSISTED_USERS = load_persisted_users()

# ====== Routes ==============================================================

@app.route("/", methods=["GET", "POST"])
def index():
    if "authenticated_users" not in session:
        session["authenticated_users"] = PERSISTED_USERS.copy()

    auth_success = session.pop("auth_success", False)
    authenticated_users = session.get("authenticated_users", {})
    active_user = session.get("active_user")

    if request.method == "POST":
        if not active_user or active_user not in authenticated_users:
            return redirect(url_for("auth"))

        customer_id = normalize_customer_id(request.form.get("customer_id", ""))
        manager_id_override = normalize_customer_id(request.form.get("manager_id", "")) or None

        auth_file = authenticated_users[active_user]
        client = load_client_with_optional_login(auth_file, manager_id_override)

        filepath = generate_google_ads_report(customer_id, client)
        session["latest_report"] = filepath
        return redirect(url_for("report"))

    return render_template(
        "index.html",
        authenticated=bool(active_user),
        auth_success=auth_success,
        authenticated_users=authenticated_users,
        active_user=active_user
    )

@app.route("/auth")
def auth():
    session.pop("oauth_state", None)
    flow = get_flow()
    auth_url, state = flow.authorization_url(
        access_type="offline",
        include_granted_scopes=False,
        prompt="consent"
    )
    session["oauth_state"] = state
    return redirect(auth_url)

@app.route("/callback")
def callback():
    stored_state = session.get("oauth_state")
    returned_state = request.args.get("state")
    if stored_state != returned_state:
        return "⚠️ OAuth state mismatch. Try <a href='/auth'>again</a>."

    flow = get_flow(state=stored_state)
    flow.fetch_token(authorization_response=request.url)
    credentials = flow.credentials

    try:
        id_info = id_token.verify_oauth2_token(credentials.id_token, requests.Request())
        email = (id_info.get("email") or "").lower()
        if not email:
            return "Could not read account email from Google."
    except Exception as e:
        return f"Failed to obtain user email from id_token: {e}"

    refresh_token = credentials.refresh_token
    if not refresh_token:
        return (
            "<h3>❌ No refresh token returned</h3>"
            "<p>Please revoke this app at "
            "<a target='_blank' href='https://myaccount.google.com/permissions'>Google Account &gt; Security &gt; Third-party access</a>, "
            "then sign in again in an incognito window.</p>"
        )

    user_yaml = user_yaml_path_for_email(email)
    existing = read_yaml(user_yaml) if os.path.exists(user_yaml) else {}

    user_config = deepcopy(base_config)
    if "login_customer_id" in existing:
        user_config["login_customer_id"] = existing["login_customer_id"]

    user_config["refresh_token"] = refresh_token
    write_yaml(user_yaml, user_config)

    session.setdefault("authenticated_users", {})
    session["authenticated_users"][email] = user_yaml
    PERSISTED_USERS[email] = user_yaml

    session["active_user"] = email
    session["auth_success"] = True
    session.pop("oauth_state", None)

    return redirect(url_for("index"))

@app.route("/switch_user/<identifier>")
def switch_user(identifier):
    authenticated_users = session.get("authenticated_users", {})
    if identifier in authenticated_users:
        session["active_user"] = identifier
        return redirect(url_for("index"))
    return "User not found.", 404

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("index"))

@app.route("/report")
def report():
    filepath = session.get("latest_report")
    if not filepath or not os.path.exists(filepath):
        return "No report found. Please generate again."
    structured_report = parse_docx_to_structured(filepath)
    return render_template(
        "report.html",
        structured_report=structured_report,
        download_link=url_for("download_file", filename=os.path.basename(filepath))
    )

@app.route("/download/<filename>")
def download_file(filename):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename, as_attachment=True)

@app.route("/section/<section_id>")
def section_detail(section_id):
    filepath = session.get("latest_report")
    if not filepath or not os.path.exists(filepath):
        return "No report found. Please generate again."

    if section_id == "heatmaps":
        return render_template("section.html", section={
            "title": "📊 Heatmaps",
            "content": [
                {"type": "image", "metric": "Clicks"},
                {"type": "image", "metric": "Conversions"},
                {"type": "image", "metric": "CVR"}
            ]
        })

    try:
        section_index = int(section_id)
        structured_report = parse_docx_to_structured(filepath)
        if section_index < 0 or section_index >= len(structured_report):
            return "Invalid section ID."
        section = structured_report[section_index]
        return render_template("section.html", section=section)
    except ValueError:
        return "Invalid section ID (not a number)."

@app.route("/report_images/<filename>")
def report_images(filename):
    return send_from_directory("report_images", filename)

# ====== Utilities (table/docx parsing) =====================================
def try_parse_to_table(text):
    if not text or not isinstance(text, str):
        return None
    cleaned = text.strip().replace("“", '"').replace("”", '"')
    try:
        data = json.loads(cleaned)
        if isinstance(data, dict):
            data = [data]
        if isinstance(data, list) and all(isinstance(d, dict) for d in data):
            headers = list(data[0].keys())
            rows = [[d.get(h, "") for h in headers] for d in data]
            return {"headers": headers, "rows": rows}
    except Exception:
        try:
            data = ast.literal_eval(cleaned)
            if isinstance(data, list) and all(isinstance(d, dict) for d in data):
                headers = list(data[0].keys())
                rows = [[d.get(h, "") for h in headers] for d in data]
                return {"headers": headers, "rows": rows}
        except Exception:
            pass

    lines = [line.strip("•*- ") for line in cleaned.splitlines() if "|" in line]
    rows = [line.split("|")[:3] for line in lines if len(line.split("|")) >= 3]
    if rows:
        return {"headers": ["Characteristic", "Insight", "Recommendation"], "rows": rows}

    for sep in [",", "\t"]:
        lines = [line for line in cleaned.splitlines() if sep in line]
        rows = [line.split(sep)[:3] for line in lines if len(line.split(sep)) >= 3]
        if rows:
            return {"headers": ["Characteristic", "Insight", "Recommendation"], "rows": rows}
    return None

def parse_docx_to_structured(path):
    from docx import Document
    doc = Document(path)
    structured = []
    current_section = {
        "title": "⭐ Introduction",
        "content": [{
            "type": "paragraph",
            "content": (
                "Hello, this is a structured report generated from a Google Ads audit document. "
                "It contains insights, visualizations, and optimization suggestions."
            )
        }]
    }
    para_index = 0
    table_index = 0

    for element in doc.element.body:
        if element.tag.endswith("}p"):
            if para_index >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[para_index]
            para_index += 1
            text = para.text.strip()
            if not text:
                continue
            if getattr(para.style, "name", "").startswith("Heading") or text.startswith("⭐"):
                if current_section["content"]:
                    structured.append(current_section)
                current_section = {"title": text, "content": []}
            else:
                table_result = try_parse_to_table(text)
                if table_result:
                    current_section["content"].append({
                        "type": "table",
                        "headers": table_result["headers"],
                        "rows": table_result["rows"]
                    })
                else:
                    current_section["content"].append({
                        "type": "paragraph",
                        "content": text
                    })
        elif element.tag.endswith("}tbl"):
            if table_index >= len(doc.tables):
                continue
            table = doc.tables[table_index]
            table_index += 1
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                current_section["content"].append({
                    "type": "table",
                    "headers": rows[0],
                    "rows": rows[1:]
                })

    if current_section["content"]:
        structured.append(current_section)
    return structured

# ====== Run ================================================================
if __name__ == "__main__":
    app.run(host='0.0.0.0', port=int(os.environ.get("PORT", 5000)))
